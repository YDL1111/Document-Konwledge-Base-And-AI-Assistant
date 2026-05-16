import sys
sys.path.insert(0, r"D:\ResumeProjects\gen-docx")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_qa(doc, q_text, a_text, high_freq=False, deep_dive=False):
    p = doc.add_paragraph()
    marker = ""
    if high_freq:
        marker = "[HIGH_FREQ]  "
    elif deep_dive:
        marker = "[DEEP_DIVE]  "
    p.add_run("Q" + marker).bold = True
    p.add_run(q_text).bold = True
    p.paragraph_format.space_after = Pt(4)
    pa = doc.add_paragraph()
    pa.paragraph_format.space_after = Pt(12)
    pa.paragraph_format.space_before = Pt(2)
    run_a = pa.add_run(a_text)
    run_a.font.size = Pt(10.5)
    run_a.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# Title Page
for _ in range(4):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("RAG 智能问答助手（Python 端）")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("项目面试 QA 全解")
r2.font.size = Pt(22)
r2.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6E)

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("技术栈：FastAPI + LangChain 0.3.x + ChromaDB + BGE-M3 + DeepSeek v4 Flash")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(3):
    doc.add_paragraph()
tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run("[HIGH_FREQ] = 面试最高频题 | [DEEP_DIVE] = 面试官最爱深挖的点")
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0x99, 0x99, 0x33)

doc.add_page_break()

# ======== Part 1: Project Overview ========
add_heading_styled(doc, "第一部分：项目概述（面试官开场题）", 1)

add_qa(doc,
    "简单介绍一下你做过的 AI 项目？",
    "我负责的是一个企业级 RAG 智能问答系统的 Python 端，定位是「AI 大脑」。整个系统采用 Java + Python 双端架构，Java 端负责文档管理、权限控制和前端交互，Python 端专注于非结构化文档解析、向量化、智能检索和大模型问答。\n\nPython 端的技术栈是 FastAPI + LangChain 0.3.x + ChromaDB，使用 BGE-M3 做 Embedding，DeepSeek v4 Flash 做大模型生成，Celery + Redis 处理异步任务。我主要实现了文档入库异步流水线、向量检索+Rerank 的智能问答接口，以及 API Key 鉴权等工程能力。",
    high_freq=True)

add_qa(doc,
    "这个项目中你负责了什么？遇到了什么困难？",
    "我独立负责 Python 端 AI 大脑的全栈实现。核心工作包括：\n\n1. 文档解析与切分：实现了支持 PDF（PyMuPDF）、Word、Excel、PPT 的多格式解析，设计了基于 RecursiveCharacterTextSplitter 的分层切分策略\n2. 向量化与检索：将 Embedding 从 Ollama 本地模型替换为 BGE-M3，在 ChromaDB 中实现了带 metadata 过滤的向量检索\n3. RAG 问答服务：用 LangChain LCEL 链式调用组装了 Prompt → LLM → 输出的完整推理管道，接入了 DeepSeek v4 Flash API\n4. 异步任务系统：把同步的文档处理改造为 Celery 异步任务 + Redis 状态管理\n\n遇到的困难主要是 BGE-M3 模型加载时的内存管理（首次加载约 2.2GB），以及 ChromaDB 在大规模 chunk 插入时的性能优化。解决方案是分批写入（每批 50 条）和增加 embedding 缓存。",
    deep_dive=True)

add_qa(doc,
    "为什么要做 RAG？为什么不用微调？",
    "选择 RAG 而不是微调主要基于三个原因：\n\n1. 数据时效性：企业文档频繁更新，微调需要重新训练模型成本高，RAG 只需重新索引文档即可，实时更新\n2. 知识溯源：RAG 天然带来源标注（chunk 来自哪个文档哪一页），微调做不到这一点。企业场景中「答案从哪来」是刚需\n3. 领域适配：BGE-M3 在中文场景下效果已经很好，不需要额外微调 Embedding 模型。生成侧用 DeepSeek v4 Flash 的通用能力配合 RAG prompt 约束，效果超过领域微调\n\n当然，如果知识库规模极小（几十份文档）且问题高度垂直，微调可能更简洁。但我们的场景是企业级多部门文档管理，RAG 是更灵活的选择。",
    high_freq=True)

# ======== Part 2: RAG Core Workflow ========
add_heading_styled(doc, "第二部分：RAG 核心链路", 1)

add_qa(doc,
    "请描述一个完整的 RAG 系统工作流程",
    "完整的 RAG 分为离线索引构建和在线查询两个阶段：\n\n【离线索引构建】\n1. 接收文档：Java 端上传文件到 MinIO，通过 HTTP 调用 Python 端的 /ingest 接口\n2. 文档解析：Python 从 MinIO 拉取文件，根据文件类型选择解析器（PDF→PyMuPDF，Word→python-docx）\n3. 文本切分：用 RecursiveCharacterTextSplitter 按 chunk_size=600、overlap=80 将文档切分成小块\n4. 向量化：BGE-M3 模型将每个 chunk 编码为 1024 维向量\n5. 存储：向量 + 原始文本 + 元数据（文档 ID、标题、页码、部门、可见范围）存入 ChromaDB\n\n【在线查询】\n1. 用户提问：Java 端转发到 Python 的 /query 接口，附带用户权限信息\n2. 查询向量化：BGE-M3 将用户问题编码为向量\n3. 向量检索：在 ChromaDB 中进行相似度搜索，带 metadata where 过滤（权限控制）\n4. Rerank 重排：BGE-Reranker 对检索结果重新打分排序\n5. Context 组装：取 top-5 作为上下文，组装 Prompt\n6. LLM 生成：DeepSeek v4 Flash 根据 Prompt 生成答案\n7. 结果返回：返回答案 + 来源标注 + 置信度",
    high_freq=True)

add_qa(doc,
    "RAG 能解决大模型的哪些根本问题？",
    "三个核心问题：\n\n1. 知识时效性：大模型训练有截止日期，之后的新知识不知道。RAG 通过外部知识库实时注入，解决了这个问题\n2. 幻觉问题：大模型会编造答案。RAG 通过「仅根据以下参考材料回答」的 prompt 约束 + 来源标注，大幅降低幻觉\n3. 领域知识缺失：通用大模型在垂直领域（如企业内部制度）知识薄弱。RAG 把企业私有知识作为上下文注入，弥补了这一点",
    high_freq=True)

add_qa(doc,
    "你的项目里 RAG 的工作流程是怎样的？",
    "我的项目严格遵循标准 RAG 架构，但有一些工程化定制：\n\n1. 多路权限控制：检索时通过 ChromaDB 的 where 过滤，只召回用户有权限的文档 chunk（如 department=''hr'' OR visibility=''public''）\n2. 版本管理：每个 chunk 携带 version 字段，文档更新时先删除旧版本 chunk 再插入新版本\n3. 异步处理：文档解析+向量化是异步的（Celery 任务），用户提交后返回 task_id，前端轮询状态\n4. 安全护栏：输入端做 prompt injection 检测，输出端校验是否包含来源标注\n5. 错误降级：检索不到结果时（rerank 最高分 < 0.3），不生成答案，直接返回「未找到相关信息」",
    deep_dive=True)

# ======== Part 3: Technology Selection ========
add_heading_styled(doc, "第三部分：技术选型与架构设计", 1)

add_qa(doc,
    "为什么选 LangChain 不选 LlamaIndex？",
    "选了 LangChain 主要基于三点：\n\n1. 生态最成熟：LangChain 是目前 RAG 领域市场占有率最高的框架，社区资源、教程、面试题目覆盖最多，学习和求职性价比最高\n2. 组件即插即用：LangChain 的 LCEL（LangChain Expression Language）让我可以很方便地替换底层组件——换 LLM 只需要改 import 和参数，换 Embedding 也同理。我的项目中从 Ollama 换到 DeepSeek + BGE-M3，LCEL 链本身几乎没动\n3. 后续扩展：项目后期要做 Agent（Function Calling），LangChain 的 LangGraph 体系已经比较成熟\n\nLlamaIndex 确实在纯 RAG 场景下 API 更简洁，但考虑到学习面更广和后续扩展性，LangChain 更适合这个项目。",
    high_freq=True)

add_qa(doc,
    "为什么用 ChromaDB 不用 Milvus？",
    "选 ChromaDB 主要是项目阶段和规模考量：\n\n1. 数据量不大：企业知识库初期 chunk 量在十万级别，ChromaDB 完全够用。Milvus 适合千万级以上\n2. 部署轻量：ChromaDB 是嵌入式存储（文件级），不需要额外部署 etcd + minio，开发调试方便\n3. 集成简单：langchain-chroma 和 ChromaDB 原生 API 都很简洁，where 过滤能满足基本的权限需求\n\n但如果项目后续规模增长到百万 chunk 以上、需要分布式部署，我会毫不犹豫地迁到 Milvus。ChromaDB 的 metadata 过滤灵活性和查询性能在大数据量下确实不如 Milvus。",
    deep_dive=True)

add_qa(doc,
    "整个系统架构是怎样的？",
    "采用 Java + Python 双端架构：\n\n【Java 端（业务底座）】\n- Spring Boot 3.x + MyBatis-Plus + MySQL 8.0\n- 负责：用户管理、RBAC 权限、文档元数据管理、文件上传到 MinIO、前端交互\n- 不碰文档解析和向量化\n\n【Python 端（AI 大脑）— 我负责的】\n- FastAPI + LangChain 0.3.x + ChromaDB\n- 负责：从 MinIO 拉文件 → 解析 → 切分 → BGE-M3 向量化 → ChromaDB 存储 → 智能问答\n\n【通信层】\n- Java ↔ Python 走内网 REST API，API Key 鉴权（X-API-Key header）\n- 文档入库走异步：Java 调 /ingest 返回 task_id，前端轮询 /ingest/{task_id} 获取状态\n- 问答接口走同步：Java 调 /query，附带用户权限上下文\n\n【基础设施】\n- MinIO（对象存储）、Redis（Celery broker + 任务状态 + 缓存）、ChromaDB（向量库）",
    high_freq=True)

add_qa(doc,
    "为什么选 DeepSeek 不选 GPT-4 或 Claude？",
    "选择 DeepSeek v4 Flash 的核心考量：\n\n1. 中文能力：DeepSeek 在中文理解、中文生成上的表现明显优于同价位的海外模型。企业 RAG 的核心场景是中文文档问答\n2. 成本：DeepSeek 的 API 价格约为 GPT-4 的 1/10，对于企业级大量文档问答场景，成本控制很重要\n3. 延迟：DeepSeek Flash 版本响应速度很快，平均延迟在 1-2 秒，适合实时问答\n4. OpenAI 兼容：API 格式完全兼容 OpenAI，用 LangChain 的 ChatOpenAI 直接对接，零适配成本\n\n如果是英文为主或需要最强推理能力的场景，我会选 Claude 或 GPT-4。但中文企业 RAG，DeepSeek 性价比最高。",
    high_freq=True)

add_qa(doc,
    "你的项目用了哪些中间件？为什么？",
    "核心中间件：\n\n1. Redis：用作 Celery 的 broker 和 result backend，同时存储任务状态（task:xxx hash）和查询缓存\n2. MinIO：企业私有化对象存储，S3 协议兼容。Java 端上传文件到 MinIO，Python 端通过 MinIO SDK 拉取，避免了文件传输的 HTTP 开销\n3. ChromaDB：轻量级向量数据库，存储 embedding 和 chunk 元数据。选它因为项目初期数据量不大，部署维护成本最低",
    deep_dive=True)

# ======== Part 4: Document Parsing ========
add_heading_styled(doc, "第四部分：文档解析与切分", 1)

add_qa(doc,
    "讲讲文档是怎么切的？Chunking 策略是什么？",
    "我用了 LangChain 的 RecursiveCharacterTextSplitter，核心参数：\n\n- chunk_size = 600（中文字符数，不是 token 数）\n- chunk_overlap = 80（相邻块重叠 80 字）\n- separators = [\"\\n\\n\", \"\\n\", \"。\", \"！\", \"？\", \";\", \" \", \"\"]\n- length_function = len（按字符数计算，适合中文）\n\n切分逻辑是：先尝试按最大的分隔符（\\n\\n）切，切出来的块如果超过 600 字，就依次用更小的分隔符（\\n、句号等）继续切，直到每块都在 600 字以内。\n\noverlap 设为 80 是为了避免语义被切断。比如一个概念的描述跨了两个块，overlap 能保证第二个块开头有上下文。",
    high_freq=True)

add_qa(doc,
    "怎么解决语义被切割的问题？",
    "三层方案：\n\n1. Overlap（重叠窗口）：chunk_overlap=80，确保相邻块有 80 字重叠，保留跨块语义\n2. 分隔符优先级：separators 从大到小排列（段落→换行→句号→空格→空字符串），优先在语义边界处切分\n3. 表格单独处理：Excel 的每个 sheet 作为独立 chunk 不拆分，Word 中的表格内容单独合并为一个 chunk\n\n如果未来需要更精细的切分，可以引入语义边界检测（基于 embedding 相似度判断句子间语义连贯性），但当前阶段 overlap + 分隔符优先级已经够用。",
    deep_dive=True)

add_qa(doc,
    "PDF 是怎么解析的？为什么用 PyMuPDF 不用 pypdf？",
    "PDF 解析用 PyMuPDF（fitz），理由：\n\n1. 中文提取质量：PyMuPDF 对中文 PDF 的文本提取明显优于 pypdf。pypdf 经常出现乱码或文字丢失，PyMuPDF 的 get_text() 基本能完整提取\n2. 页码保留：PyMuPDF 天然保留页码信息（page number），对溯源标注「[来源：《文档》第X页]」很重要\n3. 图片处理：PDF 中纯图片页可以用 page.get_pixmap() 转为图片，后续可扩展 OCR 能力\n\n解析后的每个页面都带 metadata（source=文件路径, page=页码），传给后续的 chunk 切分器。",
    deep_dive=True)

add_qa(doc,
    "Excel 和 Word 怎么处理？",
    "Excel：用 openpyxl 读取 .xlsx，每个 sheet 独立处理。遍历所有行，将非空行拼接为 ''列1 | 列2 | 列3'' 格式的文本。每个 sheet 生成一个 Document 对象。\n\nWord：用 python-docx 读取 .docx，提取所有段落文本（过滤掉空行），合并表格内容（每行 cell 用 '' | '' 拼接）。正文和表格合并后作为单个 Document 传入切分器。\n\n两者的核心原则：保留原始文档结构信息（页码/sheet 名）作为 metadata，方便溯源。")

# ======== Part 5: Embedding & Vector Search ========
add_heading_styled(doc, "第五部分：Embedding 与向量检索", 1)

add_qa(doc,
    "Embedding 是什么？为什么选 BGE-M3？",
    "Embedding 是把非结构化文本转化为固定长度数值向量的过程。转化后的向量在空间中，语义相似的文本距离更近。在 RAG 中，用户的提问被转化为 embedding 后，可以在向量数据库中通过相似度搜索找到最相关的文档片段。\n\n选 BGE-M3 的理由：\n\n1. 中文能力：BGE-M3 是智源出品的模型，在中文 MTEB 榜单上排名顶尖。相比 OpenAI text-embedding-3 等英文模型，BGE-M3 对中文语义的理解显著更好\n2. 多语言支持：1024 维向量，支持中文/英文/多语言，企业场景中英文混合文档都能处理\n3. 参数适中：2.2GB 模型大小，CPU 可推理，不需要 GPU\n4. LangChain 集成：通过 langchain-huggingface 的 HuggingFaceEmbeddings 直接对接，API 简洁",
    high_freq=True)

add_qa(doc,
    "你用的是哪种向量数据库？数据量多大？性能如何？",
    "用的是 ChromaDB 0.5.3，本地文件持久化（PersistentClient）。\n\n当前数据量级：\n- 文档数：约 100-500 份\n- chunk 数：约 5,000-20,000 条\n- 向量维度：1024（BGE-M3）\n\n性能表现：\n- 检索延迟：50-150ms（本地 SSD，数据量小时很快）\n- 插入延迟：批量写入（每批 50 条），约 1-2 秒/百条\n- 内存占用：ChromaDB + embedding 模型约 3-4GB\n\n瓶颈：ChromaDB 在 chunk 数超过 50 万后，metadata 过滤（where 条件）性能会明显下降。如果项目规模增长到那个级别，需要迁到 Milvus。\n\n当前阶段 ChromaDB 的优势是开发调试快、不需要额外服务部署。",
    deep_dive=True)

add_qa(doc,
    "做过向量数据库的对比选型吗？",
    "对比过 ChromaDB、Milvus 和 Qdrant：\n\nChromaDB：嵌入式的，Python 一行代码就能用。优势是简单、开发快；劣势是不支持分布式、metadata 过滤有限（$in、$and 部分支持）、大数据量性能下降。\n\nMilvus：企业级分布式向量数据库。优势是功能全（支持复杂 metadata 过滤、标量索引）、支持分布式部署；劣势是部署重（需要 etcd + minio）、学习成本高。\n\nQdrant：介于两者之间，Rust 编写性能好，支持 metadata 过滤。但生态不如 ChromaDB 成熟。\n\n最终选了 ChromaDB 因为：项目初期 chunk 量不大、开发阶段需要快速迭代、团队没有运维 Milvus 的经验。后续如果需要迁移，langchain-milvus 的 API 和 langchain-chroma 几乎一致，迁移成本低。",
    deep_dive=True)

add_qa(doc,
    "向量检索和关键词检索有什么区别？",
    "核心区别在匹配方式：\n\n向量检索（语义匹配）：\n- 原理：把问题和文档都转成 embedding 向量，计算余弦相似度\n- 优势：能理解语义。搜''年假几天''能匹配到''带薪年假根据工龄计算''这种不同表述的文本\n- 劣势：专有名词、精确数字匹配差。搜''2025年Q3财报''可能匹配不到\n\n关键词检索（字面匹配）：\n- 原理：基于 TF-IDF 或 BM25，直接匹配文本中的关键词\n- 优势：精确匹配好，搜特定编号/人名/术语准\n- 劣势：不理解语义。搜''年假''匹配不到''带薪休假''\n\n我的项目当前主要用向量检索，Phase 2 加了 Rerank 后进一步提升准确率。未来如果要加混合检索，可以引入 Elasticsearch 做关键词路，两路结果合并后 Rerank。")

add_qa(doc,
    "怎么在 ChromaDB 里做权限过滤？",
    "在文档入库时，每个 chunk 的 metadata 携带了 department（部门）和 visibility（可见范围）字段：\n\n入库时 metadata = {''doc_id'': ''doc_001'', ''department'': ''hr'', ''visibility'': ''public'', ''version'': ''v1.0'', ''page'': 15}\n\n检索时通过 ChromaDB 的 where 条件过滤：\n\n# 研发部用户只能看研发部门 + 公开文档\nwhere_filter = {''$or'': [ {''department'': ''engineering''}, {''visibility'': ''public''} ]}\nresults = collection.query(query_embeddings=..., where=where_filter, n_results=20)\n\nChromaDB 的 where 支持 $eq、$in、$and、$or 等操作符，基本的权限过滤够用。但如果需要更复杂的嵌套过滤（如 AND + OR 组合），ChromaDB 的限制就会暴露出来，这时候就需要 Milvus 了。",
    deep_dive=True)

# ======== Part 6: LLM & Prompt Engineering ========
add_heading_styled(doc, "第六部分：大模型调用与 Prompt 工程", 1)

add_qa(doc,
    "你的 Prompt 模板是怎么设计的？",
    "我的 Prompt 模板分为 system prompt 和 user prompt 两部分，通过 LangChain 的 ChatPromptTemplate 组装：\n\nSystem Prompt：\n''你是一个企业知识问答助手。请仅根据以下参考材料回答问题。\n\n【参考材料】\n{context}\n\n要求：\n1. 仅根据参考材料回答，不要使用你的训练数据\n2. 如果找不到相关信息，直接回答''抱歉，当前知识库中没有找到相关的内容''\n3. 回答时标注信息来源，格式为：[来源：《文档标题》第X页]\n4. 回答简洁准确，不要过度展开\n5. 涉及数据、日期、人名时，严格使用原始值''\n\nUser Prompt：\n{question}\n\n设计要点：\n- 强约束「仅根据参考材料」→ 降低幻觉\n- 明确「找不到就说找不到」→ 避免编造\n- 指定来源格式 → 方便前端渲染溯源链接\n- 限制扩展 → 减少无关内容",
    high_freq=True)

add_qa(doc,
    "temperature 参数设为 0.3 是什么意思？",
    "temperature 控制 LLM 生成文本的随机性：\n\n- temperature=0：完全确定性，每次输出相同（适合数学/代码）\n- temperature=0.3：较低随机性，在保持准确性的同时允许一定灵活性（适合 RAG 问答）\n- temperature=1.0：较高随机性，输出更多样化（适合创意写作）\n\n我设 0.3 是因为 RAG 问答需要：\n1. 准确性优先：答案必须忠实于参考材料，不能自由发挥\n2. 但也不能太死板：同一问题不同表述时，允许一定的措辞变化\n3. 0.3 是 RAG 场景的行业默认值，实测效果稳定\n\n如果问题偏事实性（如「年假多少天」），可以降低到 0.1；如果问题偏开放性（如「请总结制度要点」），可以升到 0.5。",
    high_freq=True)

add_qa(doc,
    "怎么规避 RAG 的幻觉问题？",
    "四层防线：\n\n1. Prompt 约束：system prompt 明确「仅根据参考材料回答」「找不到就说找不到」，这是最核心的防线\n2. 来源标注：要求模型在回答中标注 [来源：《文档标题》第X页]，方便人工核验\n3. 相似度阈值：检索结果通过 rerank 后，最高分 < 0.3 时不生成答案，直接返回「未找到相关信息」\n4. 输出校验：生成后检查是否包含来源标注，没有则补充或降级处理\n\n这四层从输入到输出形成了完整的幻觉防控闭环。",
    high_freq=True)

add_qa(doc,
    "什么是 CoT（Chain of Thought）？在 RAG 里能用吗？",
    "CoT（思维链）是让模型在给出最终答案前先展示推理过程的 prompt 技术。比如加一句「请一步步思考」，模型会先分析参考材料中的相关信息，再得出结论。\n\n在 RAG 中的应用：\n- 适合场景：需要多步推理的问题（如「对比两个制度有什么差异」）\n- 不适合场景：简单事实问答（如「年假几天」），CoT 会增加 token 消耗但不会提升准确率\n- 我的项目中：当前阶段没加 CoT，因为大部分企业问答是事实性问题。但如果后续遇到复杂推理场景，可以在 Prompt 中加入「请先分析参考材料中的相关信息，然后综合回答」来触发 CoT\n\n注意：CoT 会显著增加输出 token 数，成本和延迟都会翻倍。")

# ======== Part 7: Async & Engineering ========
add_heading_styled(doc, "第七部分：异步任务与工程实践", 1)

add_qa(doc,
    "为什么文档入库要做异步处理？",
    "因为文档解析 + 向量化是耗时操作：\n\n1. 大 PDF 解析可能耗时数分钟（100 页 PDF + 向量化约 3-5 分钟）\n2. 同步 HTTP 调用会超时（默认 30 秒）\n3. 用户体验差：用户上传后卡住等几分钟后才知道结果\n\n异步方案：\n1. 用户提交 → 返回 task_id → 前端轮询状态\n2. Celery 后台消费任务：MinIO 下载 → 解析 → 切分 → 向量化 → ChromaDB 存储\n3. 每个阶段更新 Redis 中的任务状态（pending → processing → completed）\n\n这样用户提交后立即收到响应，前端可以展示进度条。",
    high_freq=True)

add_qa(doc,
    "你的异步任务是怎么设计的？Celery 怎么用的？",
    "架构是 Celery + Redis：\n\nCelery 配置：\ncelery_app = Celery(''rag_tasks'', broker=''redis://localhost:6379/0'', backend=''redis://localhost:6379/0'')\n\n异步任务：\n@celery_app.task\ndef ingest_document(task_id, file_url, document_id, metadata):\n    # 1. 下载文件\n    minio_client.download(file_url, local_path)\n    # 2. 解析\n    docs = parser.parse(local_path)\n    # 3. 切分\n    chunks = splitter.split_documents(docs)\n    # 4. 向量化\n    vectors = embedding.embed_documents([c.page_content for c in chunks])\n    # 5. 存入 ChromaDB\n    vector_store.add_documents(chunks)\n    # 6. 更新状态\n    task_manager.update(task_id, status=''completed'', chunks_count=len(chunks))\n\n任务状态通过 Redis hash 存储：task:{task_id} → {status, document_id, chunks_count, created_at, message}\n\n前端轮询 GET /ingest/{task_id} 获取最新状态。",
    deep_dive=True)

add_qa(doc,
    "API Key 鉴权是怎么实现的？",
    "用 FastAPI 的 APIKeyHeader 中间件：\n\nfrom fastapi.security import APIKeyHeader\napi_key_header = APIKeyHeader(name=''X-API-Key'', auto_error=True)\n\nasync def verify_api_key(api_key: str = Security(api_key_header)):\n    if api_key != settings.API_KEY:\n        raise HTTPException(status_code=403, detail=''Invalid API Key'')\n    return api_key\n\n所有接口都加了这个依赖：\n@router.post(''/api/v1/documents/ingest'', dependencies=[Depends(verify_api_key)])\n\nAPI Key 从环境变量读取（settings.API_KEY），Java 端每个请求头带 X-API-Key。这是内网服务间的最小安全方案——不需要 OAuth2 那么重，但比裸奔安全多了。\n\n生产环境还可以加强：HMAC 签名（防篡改）、请求频率限制、IP 白名单。",
    deep_dive=True)

add_qa(doc,
    "DeepSeek API 限流怎么处理的？",
    "用 tenacity 库做指数退避重试：\n\nfrom tenacity import retry, wait_exponential, stop_after_attempt\n\n@retry(\n    wait=wait_exponential(multiplier=1, min=2, max=30),\n    stop=stop_after_attempt(3),\n    retry=lambda x: isinstance(x, Exception) and ''rate limit'' in str(x).lower(),\n)\nasync def call_llm(messages):\n    return await llm.ainvoke(messages)\n\nDeepSeek 免费版有 RPM 限制（通常 60 RPM），如果遇到 429 状态码，tenacity 会自动等 2s → 4s → 8s → ... 后重试，最多 3 次。\n\n生产环境还可以：\n- 前端加请求排队（Redis 队列）\n- 多 API Key 轮询（分散限流）\n- 监控 token 消耗，做好成本控制")

# ======== Part 8: Agent & Tools ========
add_heading_styled(doc, "第八部分：Agent 与工具调用（Phase 4 扩展）", 1)

add_qa(doc,
    "什么是 Agent？和你现在的项目有什么区别？",
    "Agent 是比单纯 RAG 问答更进一步的概念。核心区别：\n\nRAG 问答（当前项目）：\n- 用户提问 → 检索 → 生成答案\n- 单向流程，模型只做「回答」一件事\n- 不需要工具调用\n\nAgent：\n- 用户给目标 → Agent 规划 → 选择工具 → 执行 → 汇总结果\n- 可以「做事」：查知识库、预览文档、对比多份文件、生成摘要\n- 需要 Function Calling 或 MCP 协议对接工具\n\n我的项目 Phase 4 计划就是把当前 RAG 能力封装为一个工具（search_knowledge_base），让 Agent 在多种工具中选择调用。比如用户说「对比一下 2024 和 2025 年的制度差异」，Agent 会：\n1. 调用 search_knowledge_base 查 2024 年制度\n2. 调用 search_knowledge_base 查 2025 年制度\n3. 调用 summarize_document 对比差异\n4. 返回综合结果\n\n这和当前单纯的 RAG 问答（只查一次、只回答一次）有本质区别。")

add_qa(doc,
    "了解 Function Calling 吗？",
    "Function Calling 是大模型直接调用外部工具/函数的能力。原理是：\n\n1. 开发者定义工具 schema（函数名、参数、描述）\n2. 模型在生成时，可以输出「调用 XX 函数，参数是 YY」\n3. 系统执行该函数，把结果返回给模型\n4. 模型根据结果生成最终回答\n\n在我的项目中，Phase 4 计划用 DeepSeek v4 Flash 的 Function Calling 能力实现：\n- search_knowledge_base：当前 RAG 能力封装为工具\n- preview_document：触发文档预览\n- compare_documents：多文档对比\n- summarize_document：文档摘要生成\n\nFunction Calling 和 MCP 的区别：\n- Function Calling 是模型层面的能力（OpenAI 首创，现在主流模型都支持）\n- MCP 是协议层面的标准（Anthropic 提出），用于标准化工具注册和发现\n- 两者互补：FC 让模型能调工具，MCP 让工具更容易被发现和注册")

add_qa(doc,
    "你项目里的接口是怎么设计的？",
    "Python 端设计了 5 个核心接口：\n\n1. POST /api/v1/documents/ingest — 文档入库（异步）\n   - 请求：file_url, document_id, title, department, visibility, version\n   - 响应：task_id + status ''accepted''\n   - 处理：创建任务 → Celery 异步处理 → 更新 Redis 状态\n\n2. GET /api/v1/documents/ingest/{task_id} — 查询入库状态\n   - 响应：task_id, status (pending/processing/completed/failed), chunks_count\n\n3. POST /api/v1/documents/query — 智能问答\n   - 请求：question, user_id, department, visibility_level, top_k, history\n   - 响应：answer, sources[], confidence, model\n   - 处理：鉴权 → 输入安全检测 → 检索 → Rerank → Prompt → LLM → 输出校验\n\n4. DELETE /api/v1/documents/{document_id}/version/{version} — 文档删除\n   - 从 ChromaDB 中删除对应 document_id + version 的所有 chunk\n\n5. GET /api/v1/health — 健康检查\n   - 响应：status, version, chroma_connected, deepseek_available, celery_running\n\n所有接口都带 API Key 鉴权（X-API-Key header），从环境变量读取。")

# ======== Appendix ========
add_heading_styled(doc, "附录：项目技术栈速查表", 1)

# Create a summary table
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "组件"
hdr_cells[1].text = "技术选型"
hdr_cells[2].text = "版本"

data = [
    ["AI 框架", "LangChain", "0.3.9"],
    ["向量数据库", "ChromaDB", "0.5.3"],
    ["大模型", "DeepSeek v4 Flash", "API"],
    ["Embedding", "BGE-M3", "1024 维"],
    ["Rerank", "BGE-Reranker-v2-M3", "CrossEncoder"],
    ["PDF 解析", "PyMuPDF (fitz)", "1.24+"],
    ["Word 解析", "python-docx", "1.1.2"],
    ["Excel 解析", "openpyxl", "3.1.5"],
    ["PPT 解析", "python-pptx", "1.0.2"],
    ["异步任务", "Celery + Redis", "5.4+"],
    ["文件存储", "MinIO", "7.2+"],
    ["Web 框架", "FastAPI", "0.115.5"],
    ["鉴权", "API Key", "X-API-Key header"],
]

for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

# ======== Save ========
output_path = r"D:\ResumeProjects\python项目面试QA.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
