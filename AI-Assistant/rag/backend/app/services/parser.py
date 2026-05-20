"""
Document parsing service.
"""
import os
from pathlib import Path
from typing import List, Tuple

import chardet
from langchain_core.documents import Document
from loguru import logger


class DocumentParser:
    SUPPORTED_TYPES = {
        "pdf": ["pdf"],
        "word": ["doc", "docx"],
        "excel": ["xlsx"],
        "text": ["txt", "md", "csv", "json", "xml", "html", "htm"],
    }

    def get_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower().lstrip(".")
        for file_type, exts in self.SUPPORTED_TYPES.items():
            if ext in exts:
                return file_type
        return "unknown"

    def parse(self, file_path: str, filename: str = None) -> Tuple[List[Document], dict]:
        filename = filename or os.path.basename(file_path)
        ext = Path(filename).suffix.lower().lstrip(".")
        file_type = self.get_file_type(filename)

        logger.info(f"Parsing file: {filename} (type={file_type})")

        if file_type == "pdf":
            docs = self._parse_pdf(file_path)
        elif file_type == "word":
            docs = self._parse_word(file_path)
        elif file_type == "excel":
            docs = self._parse_excel(file_path)
        elif file_type == "text":
            docs = self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        total_chars = sum(len(doc.page_content) for doc in docs)
        meta = {
            "filename": filename,
            "file_type": file_type,
            "ext": ext,
            "page_count": len(docs),
            "total_chars": total_chars,
        }
        logger.info(f"Parsed {len(docs)} pages, {total_chars} chars from {filename}")
        return docs, meta

    def _parse_pdf(self, file_path: str) -> List[Document]:
        import fitz

        docs = []
        pdf_doc = fitz.open(file_path)
        try:
            for index, page in enumerate(pdf_doc):
                text = (page.get_text() or "").strip()
                if text:
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={"source": file_path, "page": index + 1},
                        )
                    )
        finally:
            pdf_doc.close()

        return docs if docs else [Document(page_content="", metadata={"source": file_path})]

    def _parse_word(self, file_path: str) -> List[Document]:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            content = "\n".join(paragraphs)
        else:
            try:
                import docx2txt

                content = docx2txt.process(file_path)
            except Exception:
                content = f"[无法解析 .doc 文件: {file_path}]"

        return [Document(page_content=content, metadata={"source": file_path})]

    def _parse_excel(self, file_path: str) -> List[Document]:
        import openpyxl

        docs = []
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        try:
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                rows = []
                for row in worksheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        rows.append(" | ".join(row_data))
                content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
                docs.append(
                    Document(
                        page_content=content,
                        metadata={"source": file_path, "sheet": sheet_name},
                    )
                )
        finally:
            workbook.close()

        return docs

    def _parse_text(self, file_path: str) -> List[Document]:
        with open(file_path, "rb") as file:
            raw = file.read()

        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"

        try:
            content = raw.decode(encoding, errors="replace")
        except Exception:
            content = raw.decode("utf-8", errors="replace")

        return [Document(page_content=content, metadata={"source": file_path})]


parser = DocumentParser()
